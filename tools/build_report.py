from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "项目开发与预期成果推进报告_甲骨里的日光缺口.docx"
SNAPSHOT = json.loads(
    (ROOT / "data" / "published-snapshot.json").read_text(encoding="utf-8")
)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5F6B73"
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
WHITE = "FFFFFF"
RISK = "9B1C1C"
INK = "172027"


def set_fonts(run, size=None, bold=None, color=None, latin="Calibri", east="Microsoft YaHei"):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), east)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tcpr = cell._tc.get_or_add_tcPr()
    tc_mar = tcpr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblpr = table._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(sum(widths_dxa)))
    tblw.set(qn("w:type"), "dxa")
    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), "120")
    tblind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(widths_dxa[idx]))
            tcw.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row):
    trpr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trpr.append(tbl_header)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_fonts(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])
    tail = paragraph.add_run(" 页")
    set_fonts(tail, size=9, color=MUTED)


def add_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids or [0]) + 1
    num_id = max(num_ids or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    numfmt = OxmlElement("w:numFmt")
    numfmt.set(qn("w:val"), "decimal")
    lvltext = OxmlElement("w:lvlText")
    lvltext.set(qn("w:val"), "%1.")
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.extend([tabs, ind, spacing])
    level.extend([start, numfmt, lvltext, suff, ppr])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_numbered(doc, text, num_id):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    ppr = p._p.get_or_add_pPr()
    numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    numpr.extend([ilvl, numid])
    ppr.append(numpr)
    run = p.add_run(text)
    set_fonts(run, size=11, color=INK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_fonts(lead, size=11, bold=True, color=INK)
        rest = p.add_run(text[len(bold_lead):])
        set_fonts(rest, size=11, color=INK)
    else:
        run = p.add_run(text)
        set_fonts(run, size=11, color=INK)
    return p


def add_callout(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    ppr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CALLOUT_FILL)
    ppr.append(shd)
    borders = OxmlElement("w:pBdr")
    left_border = OxmlElement("w:left")
    left_border.set(qn("w:val"), "single")
    left_border.set(qn("w:sz"), "18")
    left_border.set(qn("w:space"), "8")
    left_border.set(qn("w:color"), BLUE)
    borders.append(left_border)
    ppr.append(borders)
    r1 = p.add_run(f"{label}  ")
    set_fonts(r1, size=10.5, bold=True, color=BLUE)
    r2 = p.add_run(text)
    set_fonts(r2, size=10.5, color=INK)


def add_matrix_table(doc, headers, rows, widths, centered_columns=()):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
        set_cell_shading(cell, LIGHT_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            set_fonts(run, size=9.3, bold=True, color=INK)
    repeat_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = str(text)
            paragraph = cells[idx].paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if idx in centered_columns
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.10
            for run in paragraph.runs:
                set_fonts(run, size=9.1, color=INK)
    set_table_geometry(table, widths)
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for idx, size, color, before, after in (
        (1, 16, BLUE, 16, 8),
        (2, 13, BLUE, 12, 6),
        (3, 12, DARK_BLUE, 8, 4),
    ):
        style = styles[f"Heading {idx}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.clear()
    header.paragraph_format.space_after = Pt(0)
    left = header.add_run("项目开发推进报告")
    set_fonts(left, size=9, bold=True, color=MUTED)
    header.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    right = header.add_run("\tMVP 0.4")
    set_fonts(right, size=9, color=MUTED)
    add_page_field(section.footer.paragraphs[0])

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(14)
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("项目开发与预期成果推进报告")
    set_fonts(r, size=23, bold=True, color=INK)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    r = subtitle.add_run("《甲骨里的日光缺口》｜基于 Qwen 的甲骨文日食记录多模态科学传播应用")
    set_fonts(r, size=13, color=MUTED)
    for label, value in (
        ("版本：", "MVP 0.4 实施基线"),
        ("日期：", "2026 年 7 月 20 日"),
        ("依据：", "《正式参赛立项书_甲骨里的日光缺口 - 董子怡》"),
        ("当前状态：", "产品闭环首版完成；1/7 篇核心资料入库；4 类作品待人工审核"),
        ("公众快照：", SNAPSHOT["snapshotId"]),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        a = p.add_run(label)
        set_fonts(a, size=10.5, bold=True, color=INK)
        b = p.add_run(value)
        set_fonts(b, size=10.5, color=INK)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_callout(doc, "项目决策", "项目不再只是一个公众科普网页，而是由公众科普站、研究工作台、输出工作室、审核发布中心、阿里云百炼服务和可审计证据库组成的专题内容生产系统。首期严格限定 7 篇核心资料，不扩张为通用 NotebookLM 替代产品。")

    add_heading(doc, "一、产品设计与范围", 1)
    add_body(doc, "产品采用“研究端生产、审核端把关、公众端消费稳定快照”的双端结构。公众用户看到的是经过审核的科学解释、甲骨记录和有来源问答；研究人员在本机工作台完成资料导入、候选提取、作品生成、审核和发布。")
    product_rows = [
        ("公众科普站", "公众", "日食原理、商代认识、记录浏览和有来源问答"),
        ("研究工作台", "研究人员", "PDF、网页、手动文本汇入，解析、去重和候选审核"),
        ("输出工作室", "研究人员", "从选定资料生成记录表、观点对照、讲解稿和音频导览"),
        ("审核发布中心", "审核人员", "批准或驳回知识与作品，处理失效项并发布快照"),
        ("百炼服务", "系统", "检索增强问答与草稿生成，不具备自动发布权限"),
        ("可审计证据库", "系统", "保存页码、来源关系、模型、提示词版本和审核事件"),
    ]
    add_matrix_table(doc, ("组成", "主要使用者", "首期职责"), product_rows, [1800, 1800, 5760], (1,))

    add_heading(doc, "1. 首期范围", 2)
    scope_num = add_numbering(doc)
    for item in (
        "支持 PDF、网页链接和手动文本汇入；PDF按页保存定位，网页和手动文本按单元保存。",
        "支持 SHA-256 去重、候选知识提取、人工审核、作品草稿生成和稳定快照发布。",
        "首版优先完成甲骨日食记录表、学者观点与争议对照、大众讲解稿和音频导览。",
        "公众端只读取已发布快照；未审核资料、候选知识和草稿不进入公众问答。",
    ):
        add_numbered(doc, item, scope_num)
    add_heading(doc, "2. 明确不做", 2)
    add_body(doc, "本期不实现通用个人知识库、多租户账号体系、原始 PDF 公开下载、自动断代、自动学术定论或自动发布。DOCX、Markdown、扫描图片 OCR 和音视频转录只保留后续扩展接口。")

    architecture_heading = add_heading(doc, "二、系统架构与知识治理", 1)
    add_callout(doc, "主数据流", "PDF / 网页 / 手动文本 -> 私有证据库 -> 解析单元与页码 -> 候选知识 -> 人工审核 -> 发布知识 -> 作品草稿 -> 人工审核 -> 稳定公众快照 -> 公众展示与问答。")
    store_rows = [
        ("证据库", "原始资料、文本块、页码、指纹", "仅研究端", "模型输出禁止写入"),
        ("发布知识库", "释文、著录号、年代、观点、争议", "公众快照", "只收人工批准项"),
        ("作品库", "文章、记录表、图卡、音频、课件", "审核后公开", "草稿与正式作品分状态"),
        ("来源关系", "资料、页码、知识、模型、提示词、审核", "审计与失效传播", "作品不能成为证据上游"),
    ]
    add_matrix_table(doc, ("存储", "保存内容", "服务对象", "硬性规则"), store_rows, [1440, 3000, 1800, 3120], (0, 2))

    add_heading(doc, "1. 数据模型", 2)
    add_body(doc, "资料与单元：", "资料与单元：")
    add_body(doc, "source_documents 保存资料元数据、指纹、状态和解析版本；source_units 保存 PDF 页或文本单元及定位信息。")
    add_body(doc, "知识与作品：", "知识与作品：")
    add_body(doc, "knowledge_candidates 保存候选；published_knowledge 只保存批准项；artifacts 保存草稿、批准、发布和失效状态。")
    add_body(doc, "审计与发布：", "审计与发布：")
    add_body(doc, "lineage_edges 记录上游资料、页码和知识；review_events、prompt_versions 与 publish_snapshots 记录审核、提示词和发布批次。")

    add_heading(doc, "2. 防止模型循环引用", 2)
    add_body(doc, "lineage_edges.upstream_type 只允许 source、unit 和 knowledge。模型生成结果始终先进入 draft，不能作为证据节点，也不能自动重新进入证据库；只有人工审核后，作品才可进入公众作品快照。")

    add_heading(doc, "3. 变更追踪", 2)
    add_body(doc, "资料删除、内容更新或重新解析后，系统沿来源关系将依赖的候选知识、发布知识和作品标记为 stale，即“来源已更新，需要重新审核”。已经发布的旧快照保持稳定，直到审核人员显式发布新快照。")

    workflow_heading = add_heading(doc, "三、资料研究与审核工作流", 1)
    workflow_num = add_numbering(doc)
    for item in (
        "导入：选择 PDF、网页链接或手动文本，记录来源元数据和内容指纹。",
        "解析：PDF逐页提取，网页和手动文本形成可定位单元，保留解析版本。",
        "去重：内容指纹相同的资料不创建第二份有效副本。",
        "候选提取：规则或模型仅把结论写入候选表，不直接进入发布知识。",
        "知识审核：审核人员查看原文、页码和争议，批准或驳回候选。",
        "作品生成：从选定且已审核的资料生成四类首版作品草稿。",
        "作品审核：核对引用、冲突和表达边界后批准，未批准项保持私有。",
        "快照发布：汇总有效的批准知识和作品，原子写入公众稳定快照。",
    ):
        add_numbered(doc, item, workflow_num)
    add_callout(doc, "公众边界", "原始 PDF、SQLite 数据库、逐页文本、候选知识和草稿均不通过网站公开。公众模式禁用研究 UI 与研究 API；data 目录只允许稳定快照文件，旧迁移 JSON 和变形路径返回 403。")

    add_heading(doc, "四、当前实现与数据状态", 1)
    current_rows = [
        ("核心资料", "1 / 7", "52172 论文已审核，拆分为 12 个 PDF 页单元"),
        ("发布知识", str(SNAPSHOT["audit"]["publishedKnowledge"]), "3 条记录与 2 项主题知识已批准"),
        ("公众记录", str(len(SNAPSHOT["recordsMeta"]["records"])), "《合集》11480、21298、33696"),
        ("首版作品", "4 类草稿", "记录表、观点对照、大众讲解稿、音频导览均待审核"),
        ("公众作品", str(SNAPSHOT["audit"]["publishedWorks"]), "没有把未审核草稿提前发布"),
        ("公众快照", "已发布", SNAPSHOT["snapshotId"]),
    ]
    add_matrix_table(doc, ("对象", "当前状态", "说明"), current_rows, [1800, 1800, 5760], (1,))

    add_heading(doc, "1. 公众科普站", 2)
    add_body(doc, "页面按“日食科学原理—日食类型—商代认识—甲骨记录—有来源问答”组织。记录可以同时展开，显示白话释义、年代、著录、核对页码、学者观点与争议；公众页面和公众问答读取同一份快照。")
    hero_shape = doc.add_picture(str(ROOT / "assets" / "hero-eclipse.png"), width=Inches(5.5))
    hero_shape._inline.docPr.set("title", "项目主视觉")
    hero_shape._inline.docPr.set("descr", "日食与甲骨材料的传播示意图，不作为甲骨实物、拓片或学术证据")
    picture_paragraph = doc.paragraphs[-1]
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.paragraph_format.keep_with_next = True
    cap = doc.add_paragraph("图 1  项目主视觉（传播示意图，不作为甲骨实物或学术证据）")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(8)
    for run in cap.runs:
        set_fonts(run, size=9, color=MUTED)

    add_heading(doc, "2. 研究工作台与输出工作室", 2)
    add_body(doc, "研究工作台已经支持资料导入、解析、审核、来源查看、作品生成和快照发布。音频导览首版提供浏览器语音试听；正式音频文件合成不在本期范围。无百炼密钥时使用确定性的本地生成器，配置密钥后通过统一适配层调用 Qwen。")

    plan_heading = add_heading(doc, "五、开发计划与完成判据", 1)
    phase_rows = [
        ("1", "产品基线", "完成", "设计、范围、架构、计划和预期成果已形成实施基线"),
        ("2", "数据模型与来源关系", "首版完成", "SQLite、外键、来源边和迁移工具已落地"),
        ("3", "资料解析与审核", "首版完成", "PDF、URL、手动文本和候选审核已接通"),
        ("4", "百炼与检索适配", "首版完成", "统一适配器与本地保底已接通；正式密钥待回归"),
        ("5", "输出工作室", "首版完成", "4 类草稿已生成，保持待审核"),
        ("6", "审核发布与快照", "首版完成", "审核门、失效传播和原子快照已接通"),
        ("7", "公众展示与问答", "完成", "统一读取快照，私有内容不进入公众问答"),
        ("8", "质量验证", "进行中", "8 项自动测试通过；专家、用户和视觉验收待完成"),
    ]
    add_matrix_table(doc, ("阶段", "工作包", "状态", "本阶段判据"), phase_rows, [720, 2160, 1440, 5040], (0, 2))
    add_body(doc, "首期完成门槛不是“功能按钮存在”，而是 7 篇核心资料全部经过同一闭环；4 类作品完成来源核对和人工审核；公众问答通过引用准确性、冲突处理、删除同步和无依据拒答测试。")

    add_heading(doc, "六、预期成果", 1)
    outcome_rows = [
        ("产品", "公众站、研究台、输出工作室、审核发布中心", "双端职责清楚，公众端只读快照"),
        ("数据", "7 篇论文、逐页文本、审核知识、来源关系和发布历史", "任一知识或作品可追溯到资料页码"),
        ("内容", "记录表、观点对照、大众讲解稿、音频导览", "4 类作品均通过人工审核"),
        ("技术", "SQLite 模型、解析适配、百炼接口、失效传播、快照生成", "可复现、可回归、无模型回流"),
        ("赛事", "演示系统、报告、清单、PPT/PDF、讲解稿与用户测试", "公开链接和二维码指向同一发布版本"),
    ]
    add_matrix_table(doc, ("类别", "交付物", "核心验收口径"), outcome_rows, [1440, 3840, 4080], (0,))

    add_heading(doc, "1. 建议量化指标", 2)
    metrics = [
        ("来源完整率", "100%", "正式条目包含著录、出版物、PDF页或文本定位"),
        ("页码引用率", "100%", "问答和作品引用能够回到对应证据单元"),
        ("审核覆盖率", "100%", "发布知识与公众作品均有审核事件"),
        ("失效传播率", "100%", "资料变化后所有直接依赖项进入重新审核状态"),
        ("无依据拒答率", "100%", "回归集中没有来源的问题不伪造答案或引用"),
        ("用户查找完成率", "不低于 90%", "用户能找到记录的释义、著录与争议"),
    ]
    add_matrix_table(doc, ("指标", "目标", "说明"), metrics, [2160, 1680, 5520], (0, 1))
    note = doc.add_paragraph("注：用户测试指标是下一阶段目标，不是当前已经取得的数据。")
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(4)
    for run in note.runs:
        set_fonts(run, size=9, color=RISK)

    verification_heading = add_heading(doc, "七、验证结果、风险与下一步", 1)
    add_heading(doc, "1. 已完成验证", 2)
    verify_num = add_numbering(doc)
    for item in (
        "重复资料识别、资料审核门和候选批准门通过。",
        "模型输出不回流证据库，作品创建后默认保持 draft。",
        "资料重解析和删除会使依赖知识与作品失效。",
        "公众快照不包含私有路径、原始资料或全文块。",
        "私网 URL 导入被拒绝，冲突观点分开保留。",
        "无依据问题明确拒答并返回 0 条引用。",
        "公众模式禁用研究 UI、研究 API 和私有目录。",
        "旧迁移 JSON、目录回退和变形路径在公众模式下返回 403。",
    ):
        add_numbered(doc, item, verify_num)

    add_heading(doc, "2. 当前风险", 2)
    risk_rows = [
        ("资料覆盖不足", "高", "仅 1/7 篇论文完成闭环；不能宣称记录已经穷尽"),
        ("专家复核未完成", "高", "释文、分期和候选日食对应仍需古文字学与天文学复核"),
        ("百炼正式回归未完成", "中", "适配接口已完成，但尚缺正式密钥下的系统问题集验证"),
        ("作品尚未发布", "中", "4 类首版输出保持草稿，需人工核对后才能进入作品库"),
        ("视觉与用户验收不足", "中", "自动浏览器受本地安全策略限制，仍需人工桌面/移动检查和目标用户试测"),
    ]
    add_matrix_table(doc, ("风险", "等级", "控制措施"), risk_rows, [2160, 1200, 6000], (1,))

    add_heading(doc, "3. 下一阶段顺序", 2)
    next_num = add_numbering(doc)
    for item in (
        "取得其余 6 篇核心资料，逐篇完成导入、解析、候选审核和来源登记。",
        "由古文字学与天文学指导教师复核释文、分期、日食对应与表达边界。",
        "审核 4 类首版作品，发布首个包含公众作品的稳定快照。",
        "使用正式百炼密钥运行 20 至 30 个问题的引用、冲突和拒答回归集。",
        "组织 5 至 10 名目标用户试测，并完成人工桌面与移动端视觉验收。",
        "部署公众模式，制作二维码和不超过 20 页的参赛 PPT/PDF。",
    ):
        add_numbered(doc, item, next_num)
    add_callout(doc, "优先级", "下一步先补齐 7 篇论文的材料覆盖和专家复核，再扩展更多生成类型。当前首版已经证明产品闭环可运行，决定参赛可信度的是证据质量、审核记录和真实用户验证。")

    core = doc.core_properties
    core.title = "《甲骨里的日光缺口》项目开发与预期成果推进报告"
    core.subject = "挑战杯参赛项目 MVP 0.4 产品设计、架构、开发与验收"
    core.author = "项目团队"
    core.keywords = "甲骨文, 日食, Qwen, 科学传播, MVP"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
